import os
import re
from datetime import datetime
from docx import Document

TEMPLATE_PATH = 'Affidavit_Template.docx'

def format_id(id_line):
    id_line = id_line.strip()
    id_line_lower = id_line.lower()

    aadhaar_raw = re.findall(r'\d{12}', id_line.replace(' ', ''))
    if 'aadhar' in id_line_lower and aadhaar_raw:
        aadhaar = aadhaar_raw[0]
        formatted = ' '.join([aadhaar[i:i+4] for i in range(0, 12, 4)])
        return f"{formatted} (Aadhar Card)"

    pan_match = re.findall(r'[A-Z]{5}[0-9]{4}[A-Z]', id_line.upper())
    if 'pan' in id_line_lower and pan_match:
        return f"{pan_match[0]} (PAN Card)"

    voter_id_match = re.findall(r'\b[A-Z0-9]{6}\b', id_line.upper())
    if 'voter' in id_line_lower and voter_id_match:
        return f"{voter_id_match[0]} (Voter ID Card)"

    return id_line

def get_relation_prefix(gender, relation_type):
    gender = gender.lower()
    relation_type = relation_type.lower()
    if gender == "male":
        return "S/o" if relation_type == "father" else "H/o"
    elif gender == "female":
        return "D/o" if relation_type == "father" else "W/o"
    return "S/o"

def title_case(text):
    return text.title()

def extract_data(text):
    fields = {
        "name": "", "relation": "", "relation_type": "", "age": "", "gender": "",
        "current_address": "", "id": "", "permanent_address": "", "mobile": ""
    }

    for line in text.splitlines():
        if ':' in line:
            key, val = line.split(':', 1)
            key, val = key.lower().strip(), val.strip()

            if 'name' in key and 'father' not in key and 'husband' not in key:
                fields["name"] = val.upper()
            elif 'father' in key:
                fields["relation"] = title_case(val)
                fields["relation_type"] = "father"
            elif 'husband' in key:
                fields["relation"] = title_case(val)
                fields["relation_type"] = "husband"
            elif 'age' in key:
                fields["age"] = val
            elif 'gender' in key:
                fields["gender"] = val
            elif 'current address' in key:
                fields["current_address"] = title_case(val)
            elif 'permanent address' in key:
                fields["permanent_address"] = title_case(val)
            elif 'pancard' in key or 'aadhar' in key or 'voter' in key:
                fields["id"] = format_id(line)
            elif 'mobile' in key:
                fields["mobile"] = val

    return fields

def replace_placeholders(doc, mapping):
    for para in doc.paragraphs:
        for key, value in mapping.items():
            if f"{{{{{key}}}}}" in para.text:
                inline = para.runs
                for i in range(len(inline)):
                    if f"{{{{{key}}}}}" in inline[i].text:
                        inline[i].text = inline[i].text.replace(f"{{{{{key}}}}}", value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders(cell, mapping)

def process_affidavit_text(message, date_str):
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template file '{TEMPLATE_PATH}' not found.")

    joining_date = datetime.strptime(date_str, "%Y-%m-%d")
    day = joining_date.strftime("%d")
    month = joining_date.strftime("%B").upper()
    year = joining_date.strftime("%Y")
    date_formatted = f"{day} of {month} {year}"

    data = extract_data(message)
    relation_prefix = get_relation_prefix(data["gender"], data["relation_type"])

    mapping = {
        "NAME": data["name"],
        "RELATION_PREFIX": relation_prefix,
        "RELATION_NAME": data["relation"],
        "AGE": data["age"],
        "CURRENT_ADDRESS": data["current_address"],
        "ID_PROOF": data["id"],
        "PERMANENT_ADDRESS": data["permanent_address"],
        "DATE": date_formatted
    }

    doc = Document(TEMPLATE_PATH)
    replace_placeholders(doc, mapping)

    output_dir = joining_date.strftime("%Y-%m-%d")
    os.makedirs(output_dir, exist_ok=True)

    filename = f"{data['name'].replace(' ', '_')}_Affidavit.docx"
    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)

    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return full_text, file_path
